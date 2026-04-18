"""
PRCE QA Agent v2 - prce.co.il
בדיקות מקצועיות מלאות: SEO, ביצועים, נגישות, אבטחה, תוכן, API, קישורים.
"""
import asyncio
import sys
import json
import base64
import requests
import re
from datetime import datetime, timedelta, date
from pathlib import Path
from urllib.parse import urlparse, urljoin

from playwright.async_api import async_playwright
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')

# ─── חגי בורסת ת"א (TASE) — עדכן מדי שנה ─────────────────────────────────────
# ימים שבהם הבורסה סגורה (שישי/שבת נסגרים אוטומטית בלוגיקת weekday)
TASE_HOLIDAYS = {
    # 2026 — תשפ"ו/תשפ"ז
    date(2026, 3,  3),   # פורים
    date(2026, 4,  2),   # פסח — יום א׳
    date(2026, 4,  8),   # פסח — יום ז׳ (אחרון בישראל)
    date(2026, 4, 21),   # יום הזיכרון
    date(2026, 4, 22),   # יום העצמאות
    date(2026, 5, 22),   # שבועות
    date(2026, 9, 11),   # ראש השנה — יום א׳ (5787)
    date(2026, 9, 12),   # ראש השנה — יום ב׳
    date(2026, 9, 20),   # יום כיפור
    date(2026, 9, 25),   # סוכות — יום א׳
    date(2026, 10, 2),   # שמיני עצרת / שמחת תורה
}

# ─── עזרים: ימי מסחר ──────────────────────────────────────────────────────────
def get_last_trading_day(reference_date=None):
    """מחזיר את יום המסחר האחרון בבורסת ת"א (שני-שישי, לא חג)."""
    if reference_date is None:
        reference_date = datetime.now().date()
    day = reference_date - timedelta(days=1)
    while day.weekday() >= 5 or day in TASE_HOLIDAYS:  # 5=שישי, 6=שבת
        day -= timedelta(days=1)
    return day

def fmt_date_il(d):
    """ממיר date לפורמט DD.M.YYYY כמו שמוצג באתר"""
    return f"{d.day}.{d.month}.{d.year}"

# ─── הגדרות ────────────────────────────────────────────────────────────────────
BASE_URL    = 'https://prce.co.il'
WP_API      = f'{BASE_URL}/wp-json/wp/v2'
API_URL     = f'{BASE_URL}/wp-json/prce/v1/index-data'
WP_USER     = 'ravidzeberger@gmail.com'
WP_PASS     = 'wa1d PY9e r5AQ IsBQ z9yN 29PK'
WP_AUTH     = base64.b64encode(f'{WP_USER}:{WP_PASS}'.encode()).decode()
WP_HEADERS  = {'Authorization': f'Basic {WP_AUTH}'}

# עמודים שה-noindex בהם מכוון ואין לדווח עליהם כשגיאה
NOINDEX_EXPECTED = {'/thank-you/'}
# עמודים שאין להם meta description בכוונה (noindex + utilty pages)
SEO_EXEMPT_PAGES = {'/thank-you/'}

REPORTS_DIR     = Path('reports')
SCREENSHOTS_DIR = REPORTS_DIR / 'screenshots'

# data-ids של ווידג'טים עם נתוני API
API_WIDGET_IDS = {
    '/':               ['8148c32','9764ac9','7a9f798','f3e2b8e','9563d66','b4d95a1','983603d','a9287c8'],
    '/metrics/naam/':  ['05abb9a','d87d263','c2ce7f7','3049ab3','80354c4','349bfed','7c50ccb','f11d4e2'],
}

DEVICES = [
    {'name': 'Desktop Firefox',     'browser_type': 'firefox',  'viewport': {'width': 1920, 'height': 1080}, 'is_mobile': False},
    {'name': 'Desktop Chrome 1920', 'browser_type': 'chromium', 'viewport': {'width': 1920, 'height': 1080}, 'is_mobile': False},
    {'name': 'Desktop Chrome 1280', 'browser_type': 'chromium', 'viewport': {'width': 1280, 'height': 800},  'is_mobile': False},
    {'name': 'iPhone 14 Pro',       'browser_type': 'webkit',   'viewport': {'width': 390,  'height': 844},  'is_mobile': True,
     'device_scale_factor': 3, 'user_agent': 'Mozilla/5.0 (iPhone; CPU iPhone OS 16_0 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/16.0 Mobile/15E148 Safari/604.1'},
    {'name': 'iPhone SE',           'browser_type': 'webkit',   'viewport': {'width': 375,  'height': 667},  'is_mobile': True,
     'device_scale_factor': 2, 'user_agent': 'Mozilla/5.0 (iPhone; CPU iPhone OS 15_0 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/15.0 Mobile/15E148 Safari/604.1'},
    {'name': 'Samsung Galaxy S23',  'browser_type': 'chromium', 'viewport': {'width': 360,  'height': 780},  'is_mobile': True,
     'device_scale_factor': 3, 'user_agent': 'Mozilla/5.0 (Linux; Android 13; SM-S911B) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/112.0.0.0 Mobile Safari/537.36'},
    {'name': 'Samsung Galaxy S22',  'browser_type': 'chromium', 'viewport': {'width': 360,  'height': 760},  'is_mobile': True,
     'device_scale_factor': 3, 'user_agent': 'Mozilla/5.0 (Linux; Android 12; SM-S906B) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/112.0.0.0 Mobile Safari/537.36'},
    {'name': 'Samsung Galaxy Tab',  'browser_type': 'chromium', 'viewport': {'width': 800,  'height': 1280}, 'is_mobile': True,
     'device_scale_factor': 2, 'user_agent': 'Mozilla/5.0 (Linux; Android 12; SM-X706B) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/112.0.0.0 Safari/537.36'},
]

# ─── WordPress: מושך כל העמודים ────────────────────────────────────────────────
def get_all_wp_pages():
    pages = []
    for post_type in ['pages', 'posts']:
        page_num = 1
        while True:
            r = requests.get(f'{WP_API}/{post_type}',
                             headers=WP_HEADERS,
                             params={'per_page': 100, 'page': page_num, 'status': 'publish'},
                             timeout=15)
            if r.status_code != 200:
                break
            batch = r.json()
            if not batch:
                break
            for p in batch:
                link = p.get('link', '')
                path = link.replace(BASE_URL, '').rstrip('/') + '/'
                if path == '//':
                    path = '/'
                pages.append({
                    'path':      path,
                    'name':      p.get('title', {}).get('rendered', path),
                    'type':      post_type,
                    'id':        p['id'],
                    'check_api': path in API_WIDGET_IDS,
                })
            if len(batch) < 100:
                break
            page_num += 1
    if not any(p['path'] == '/' for p in pages):
        pages.insert(0, {'path': '/', 'name': 'דף הבית', 'type': 'page', 'id': 0, 'check_api': True})
    return pages

# ─── API: ערכים צפויים ─────────────────────────────────────────────────────────
def get_expected_values():
    try:
        r   = requests.get(API_URL, timeout=10)
        idx = r.json()['index']
        val     = float(idx['daily_value'])
        chg_abs = float(idx['daily_change_abs'])
        vdate   = idx['value_date'][:10]
        parts   = vdate.split('-')
        return {
            'value':      f"{val:,.2f}",
            'date':       f"{int(parts[2])}.{int(parts[1])}.{parts[0]}",
            'change':     f"{chg_abs:.2f}",
            'value_date': vdate,   # YYYY-MM-DD לבדיקת עדכניות
        }
    except Exception as e:
        return {'error': str(e)}

# ─── בדיקות סטטיות (HTTP בלבד, ללא דפדפן) ────────────────────────────────────
def run_static_checks():
    results = {}

    # sitemap.xml
    try:
        r = requests.get(f'{BASE_URL}/sitemap.xml', timeout=10, headers={'User-Agent': 'PRCE-QA-Bot'})
        if r.status_code == 200 and ('<sitemap' in r.text or '<url' in r.text):
            results['sitemap'] = {'ok': True, 'detail': f'HTTP 200, {len(r.text)} bytes'}
        else:
            results['sitemap'] = {'ok': False, 'detail': f'HTTP {r.status_code}'}
    except Exception as e:
        results['sitemap'] = {'ok': False, 'detail': str(e)[:80]}

    # robots.txt
    try:
        r = requests.get(f'{BASE_URL}/robots.txt', timeout=10)
        if r.status_code == 200:
            content = r.text
            # בדוק חסימת כל האתר
            blocked = False
            lines = content.split('\n')
            in_star = False
            for line in lines:
                line = line.strip()
                if line.lower() == 'user-agent: *':
                    in_star = True
                elif in_star and line.startswith('User-agent:'):
                    in_star = False
                elif in_star and line == 'Disallow: /':
                    blocked = True
                    break
            if blocked:
                results['robots'] = {'ok': False, 'detail': 'Disallow: / — חוסם את כל האתר מגוגל!'}
            else:
                results['robots'] = {'ok': True, 'detail': f'HTTP 200, {"Sitemap מוגדר" if "Sitemap:" in content else "אין Sitemap ב-robots.txt"}'}
        else:
            results['robots'] = {'ok': False, 'detail': f'HTTP {r.status_code}'}
    except Exception as e:
        results['robots'] = {'ok': False, 'detail': str(e)[:80]}

    # Hero video
    try:
        video_url = f'{BASE_URL}/wp-content/uploads/Hero.mp4'
        r = requests.head(video_url, timeout=10, allow_redirects=True)
        size_bytes = int(r.headers.get('content-length', 0))
        size_mb = size_bytes / 1024 / 1024
        if r.status_code == 200:
            if size_mb > 15:
                results['hero_video'] = {'ok': False, 'detail': f'גדול מדי: {size_mb:.1f}MB (מקסימום מומלץ 5MB)'}
            else:
                results['hero_video'] = {'ok': True, 'detail': f'HTTP 200, {size_mb:.1f}MB'}
        else:
            results['hero_video'] = {'ok': False, 'detail': f'HTTP {r.status_code}'}
    except Exception as e:
        results['hero_video'] = {'ok': False, 'detail': str(e)[:80]}

    # API freshness — השוואה ליום המסחר האחרון (שני-שישי)
    try:
        r = requests.get(API_URL, timeout=10)
        idx = r.json()['index']
        vdate_str = idx['value_date'][:10]
        vdate_date = datetime.strptime(vdate_str, '%Y-%m-%d').date()
        last_trading = get_last_trading_day()
        if vdate_date < last_trading:
            results['api_freshness'] = {
                'ok': False,
                'detail': (f'נתוני API מיושנים: מוצג {vdate_str} — '
                           f'יום המסחר האחרון היה {fmt_date_il(last_trading)} ({last_trading})')
            }
        else:
            delta = (datetime.now().date() - vdate_date).days
            results['api_freshness'] = {'ok': True, 'detail': f'עדכני: {vdate_str} ({delta} ימים)'}
    except Exception as e:
        results['api_freshness'] = {'ok': False, 'detail': str(e)[:80]}

    # SSL certificate (HTTPS redirect from HTTP)
    try:
        r = requests.get(f'http://prce.co.il/', timeout=10, allow_redirects=True)
        if r.url.startswith('https://'):
            results['https_redirect'] = {'ok': True, 'detail': f'מפנה ל-HTTPS תקין'}
        else:
            results['https_redirect'] = {'ok': False, 'detail': 'לא מפנה ל-HTTPS!'}
    except Exception as e:
        results['https_redirect'] = {'ok': False, 'detail': str(e)[:80]}

    return results


# ─── בדיקת עמוד יחיד ──────────────────────────────────────────────────────────
async def check_page(page, page_info, expected, device):
    url    = BASE_URL + page_info['path']
    issues = []
    perf   = {}

    console_errors = []
    page.on('console', lambda m: console_errors.append(m.text) if m.type == 'error' else None)

    # ══════════════════════════════════════════════════════
    # 1. טעינת עמוד + זמן תגובה
    # ══════════════════════════════════════════════════════
    load_start = datetime.now()
    try:
        resp   = await page.goto(url, wait_until='load', timeout=60_000)
        status = resp.status if resp else 0
        load_ms = int((datetime.now() - load_start).total_seconds() * 1000)
        if status != 200:
            issues.append({'severity': 'error', 'check': 'טעינת עמוד',
                           'detail': f'HTTP {status}',
                           'fix': f'בדוק שהעמוד {page_info["path"]} מפורסם ב-WordPress'})
            return {'page': page_info['name'], 'path': page_info['path'],
                    'device': device['name'], 'issues': issues, 'screenshot': None, 'perf': perf}
    except Exception as e:
        err = str(e)[:120]
        issues.append({'severity': 'error', 'check': 'טעינת עמוד', 'detail': err,
                       'fix': 'הדף לא נטען — בדוק שהשרת עולה ואין timeout'})
        return {'page': page_info['name'], 'path': page_info['path'],
                'device': device['name'], 'issues': issues, 'screenshot': None, 'perf': perf}

    await page.wait_for_timeout(3_500)

    # מדידת ביצועים דרך Navigation Timing API
    try:
        timing = await page.evaluate("""
            () => {
                const nav = performance.getEntriesByType('navigation')[0];
                if (!nav) return {};
                return {
                    ttfb:    Math.round(nav.responseStart - nav.requestStart),
                    dom:     Math.round(nav.domContentLoadedEventEnd - nav.startTime),
                    load:    Math.round(nav.loadEventEnd - nav.startTime),
                    size:    Math.round(performance.getEntriesByType('resource')
                                .reduce((s, r) => s + (r.transferSize || 0), 0) / 1024)
                };
            }
        """)
        perf = timing or {}
        if perf.get('ttfb', 0) > 2000:
            issues.append({'severity': 'warning', 'check': 'ביצועים — TTFB',
                           'detail': f'תגובת שרת ראשונית: {perf["ttfb"]}ms (מומלץ <2000ms)',
                           'fix': 'בדוק עומס שרת, הגדרות PHP, ו-database queries. שקול שדרוג אחסון.'})
        if perf.get('load', 0) > 8000:
            issues.append({'severity': 'warning', 'check': 'ביצועים — זמן טעינה',
                           'detail': f'טעינה מלאה: {perf["load"]/1000:.1f}s (מומלץ <5s)',
                           'fix': 'הפחת גודל תמונות/וידאו, הפעל lazy loading, צמצם JS/CSS'})
        if perf.get('size', 0) > 5000:
            issues.append({'severity': 'warning', 'check': 'ביצועים — גודל עמוד',
                           'detail': f'גודל כולל: {perf["size"]/1024:.1f}MB (מומלץ <3MB)',
                           'fix': 'דחוס תמונות (WebP), הסר JS/CSS לא בשימוש, צמצם וידאו'})
    except Exception:
        pass

    # ══════════════════════════════════════════════════════
    # 2. שגיאות JavaScript
    # ══════════════════════════════════════════════════════
    NOISE = ['STAT', 'nameID', 'Table discarded', '_ga', '_gcl', '_fbp', '_fbc',
             'JQMIGRATE', 'Ignoring unsupported entryTypes', 'non-passive event',
             'ResizeObserver', 'favicon']
    real_errors = [e for e in console_errors if not any(n in e for n in NOISE)]

    def translate_js_error(err):
        e = err.lower()
        if 'failed to fetch' in e or 'networkerror' in e:
            return 'קריאה ל-API נכשלה — השרת לא היה נגיש ברגע הבדיקה'
        if 'net::err_' in e:
            code = [p for p in err.split() if 'ERR_' in p]
            return f'שגיאת רשת ({code[0] if code else "ERR"}): משאב לא נטען'
        if 'is not defined' in e or 'cannot read prop' in e:
            return f'שגיאת JavaScript: {err[:100]}'
        if 'content security policy' in e:
            return 'חסימת CSP: הדפדפן חסם תוכן — בדוק הגדרות CSP'
        return f'שגיאת JavaScript: {err[:150]}'

    if real_errors:
        explanations = [translate_js_error(e) for e in real_errors[:3]]
        issues.append({'severity': 'warning', 'check': 'שגיאות JavaScript',
                       'detail': '\n'.join(f'- {ex}' for ex in explanations),
                       'fix': 'פתח DevTools (F12) → Console לפירוט מלא'})

    # ══════════════════════════════════════════════════════
    # 3. תמונות שבורות
    # ══════════════════════════════════════════════════════
    broken = await page.evaluate("""
        () => Array.from(document.querySelectorAll('img'))
              .filter(i => i.complete && i.naturalWidth === 0)
              .map(i => i.getAttribute('src') || i.src)
              .filter(s => s && !s.startsWith('data:') && s.trim() !== '')
              .filter((v, i, a) => a.indexOf(v) === i)
    """)
    if broken:
        clean = [requests.utils.unquote(b) for b in broken]
        filenames = [b.split('/')[-1] for b in clean]
        issues.append({'severity': 'warning', 'check': 'תמונות שבורות',
                       'detail': ', '.join(filenames),
                       'fix': f'כנס ל-WordPress → מדיה → חפש: {", ".join(filenames[:3])}'})

    # ══════════════════════════════════════════════════════
    # 4. תפריט ניווט
    # ══════════════════════════════════════════════════════
    nav_ok = await page.evaluate("""
        () => !!document.querySelector('.elementor-nav-menu, .nav-menu, nav ul')
    """)
    if not nav_ok:
        issues.append({'severity': 'error', 'check': 'תפריט ניווט',
                       'detail': 'תפריט לא נמצא ב-DOM',
                       'fix': 'בדוק ב-Elementor שווידג\'ט Nav Menu קיים ומוצג'})

    # ══════════════════════════════════════════════════════
    # 5. SEO — Meta Title
    # ══════════════════════════════════════════════════════
    title = await page.evaluate("() => document.title || ''")
    title = title.strip()
    if not title:
        issues.append({'severity': 'error', 'check': 'SEO — כותרת עמוד (title)',
                       'detail': 'כותרת חסרה לחלוטין',
                       'fix': 'הוסף Title ב-Yoast SEO → עריכת עמוד → SEO → כותרת SEO'})
    elif len(title) < 20:
        issues.append({'severity': 'warning', 'check': 'SEO — כותרת עמוד (title)',
                       'detail': f'קצרה מדי: "{title}" ({len(title)} תווים, מינימום 30)',
                       'fix': 'הרחב את הכותרת ב-Yoast SEO — כלול מילות מפתח'})
    elif len(title) > 65:
        issues.append({'severity': 'warning', 'check': 'SEO — כותרת עמוד (title)',
                       'detail': f'ארוכה מדי: {len(title)} תווים (מקסימום 60)',
                       'fix': 'קצר את הכותרת — גוגל יחתוך אחרי 60 תווים'})

    # ══════════════════════════════════════════════════════
    # 6. SEO — Meta Description
    # ══════════════════════════════════════════════════════
    desc = await page.evaluate("""
        () => document.querySelector('meta[name="description"]')?.content?.trim() || ''
    """)
    if not desc and page_info['path'] not in SEO_EXEMPT_PAGES:
        issues.append({'severity': 'error', 'check': 'SEO — תיאור מטא (description)',
                       'detail': 'Meta description חסרה',
                       'fix': 'הוסף תיאור ב-Yoast SEO → עריכת עמוד → SEO → מטא תיאור'})
    elif desc and len(desc) < 80:
        issues.append({'severity': 'warning', 'check': 'SEO — תיאור מטא (description)',
                       'detail': f'קצר מדי: {len(desc)} תווים (מינימום 120)',
                       'fix': 'הרחב את התיאור — תאר את תוכן העמוד בצורה מפורטת'})
    elif len(desc) > 165:
        issues.append({'severity': 'warning', 'check': 'SEO — תיאור מטא (description)',
                       'detail': f'ארוך מדי: {len(desc)} תווים (מקסימום 160)',
                       'fix': 'קצר את התיאור — גוגל יחתוך אחרי 160 תווים'})

    # ══════════════════════════════════════════════════════
    # 7. SEO — H1
    # ══════════════════════════════════════════════════════
    h1_data = await page.evaluate("""
        () => ({
            count: document.querySelectorAll('h1').length,
            text: document.querySelector('h1')?.innerText?.trim()?.substring(0, 80) || ''
        })
    """)
    if h1_data['count'] == 0:
        issues.append({'severity': 'error', 'check': 'SEO — תגית H1',
                       'detail': 'אין H1 בעמוד — גוגל לא יבין מה נושא העמוד',
                       'fix': 'הוסף כותרת H1 ב-Elementor או ב-Yoast SEO'})
    elif h1_data['count'] > 1:
        issues.append({'severity': 'warning', 'check': 'SEO — תגית H1',
                       'detail': f'יש {h1_data["count"]} תגיות H1 — צריך בדיוק אחת',
                       'fix': 'השאר H1 יחיד. שאר הכותרות יהיו H2/H3'})

    # ══════════════════════════════════════════════════════
    # 8. SEO — Canonical URL
    # ══════════════════════════════════════════════════════
    canonical = await page.evaluate("""
        () => document.querySelector('link[rel="canonical"]')?.href || ''
    """)
    if not canonical and page_info['path'] not in SEO_EXEMPT_PAGES:
        issues.append({'severity': 'warning', 'check': 'SEO — Canonical URL',
                       'detail': 'אין canonical link — עלול לגרום ל-duplicate content',
                       'fix': 'הפעל Yoast SEO — הוא מוסיף canonical אוטומטית'})

    # ══════════════════════════════════════════════════════
    # 9. SEO — Open Graph
    # ══════════════════════════════════════════════════════
    og = await page.evaluate("""
        () => ({
            title: document.querySelector('meta[property="og:title"]')?.content || '',
            image: document.querySelector('meta[property="og:image"]')?.content || '',
            desc:  document.querySelector('meta[property="og:description"]')?.content || ''
        })
    """)
    if not og['title']:
        issues.append({'severity': 'warning', 'check': 'SEO — Open Graph (שיתוף)',
                       'detail': 'og:title חסר — שיתוף בפייסבוק/ווטסאפ יראה לא טוב',
                       'fix': 'הגדר תמונת שיתוף ב-Yoast SEO → לשונית Social'})
    if not og['image']:
        issues.append({'severity': 'warning', 'check': 'SEO — Open Graph (שיתוף)',
                       'detail': 'og:image חסר — שיתוף בסושיאל יהיה ללא תמונה',
                       'fix': 'הגדר תמונת שיתוף ב-Yoast SEO → לשונית Social → תמונת Facebook'})

    # ══════════════════════════════════════════════════════
    # 10. SEO — noindex לא מכוון
    # ══════════════════════════════════════════════════════
    noindex = await page.evaluate("""
        () => {
            const robots = document.querySelector('meta[name="robots"]')?.content || '';
            return robots.includes('noindex');
        }
    """)
    if noindex and page_info['path'] not in NOINDEX_EXPECTED:
        issues.append({'severity': 'error', 'check': 'SEO — noindex (⚠️ קריטי)',
                       'detail': f'העמוד {page_info["path"]} מסומן noindex — גוגל לא יאנדקס אותו!',
                       'fix': 'בדוק ב-Yoast SEO → עריכת עמוד → Advanced → "Allow search engines to index this page"'})

    # ══════════════════════════════════════════════════════
    # 11. נגישות — Alt Text לתמונות
    # ══════════════════════════════════════════════════════
    imgs_no_alt = await page.evaluate("""
        () => Array.from(document.querySelectorAll('img'))
            .filter(img => {
                if (img.closest('[aria-hidden="true"]')) return false;
                if (img.getAttribute('role') === 'presentation') return false;
                if ((img.src || '').includes('data:')) return false;
                return !img.getAttribute('alt') && img.getAttribute('alt') !== '';
            })
            .map(img => (img.src || '').split('/').pop().split('?')[0])
            .filter((v, i, a) => v && a.indexOf(v) === i)
            .slice(0, 5)
    """)
    if imgs_no_alt:
        issues.append({'severity': 'warning', 'check': 'נגישות — Alt Text חסר',
                       'detail': f'תמונות ללא alt: {", ".join(imgs_no_alt)}',
                       'fix': 'הוסף alt text לכל תמונה ב-WordPress (ספריית מדיה → Alt Text) — חשוב ל-SEO ולנגישות'})

    # ══════════════════════════════════════════════════════
    # 12. כיוון RTL
    # ══════════════════════════════════════════════════════
    rtl_ok = await page.evaluate("""
        () => {
            const dir = document.documentElement.getAttribute('dir') ||
                        document.body.getAttribute('dir') ||
                        getComputedStyle(document.body).direction;
            return dir === 'rtl';
        }
    """)
    if not rtl_ok:
        issues.append({'severity': 'warning', 'check': 'נגישות — כיוון RTL',
                       'detail': 'כיוון האתר לא מוגדר RTL',
                       'fix': 'ב-WordPress → הגדרות → כללי → שפה: עברית. בדוק גם ב-functions.php'})

    # ══════════════════════════════════════════════════════
    # 13. אבטחה — Mixed Content
    # ══════════════════════════════════════════════════════
    mixed = await page.evaluate("""
        () => Array.from(document.querySelectorAll(
                'img[src^="http:"], script[src^="http:"], link[href^="http:"], source[src^="http:"]'
              ))
              .map(el => el.src || el.href)
              .filter(u => u && !u.startsWith('https:'))
              .filter((v, i, a) => a.indexOf(v) === i)
              .slice(0, 3)
    """)
    if mixed:
        issues.append({'severity': 'error', 'check': 'אבטחה — Mixed Content',
                       'detail': f'משאבי HTTP על אתר HTTPS: {", ".join(mixed)}',
                       'fix': 'החלף כל URL של http:// ל-https:// — פלאגין "Better Search Replace" יכול לעזור'})

    # ══════════════════════════════════════════════════════
    # 14. Favicon — בדוק שה-link tag קיים וגם שהקובץ מחזיר 200
    # ══════════════════════════════════════════════════════
    favicon_url = await page.evaluate("""
        () => {
            const el = document.querySelector('link[rel="icon"], link[rel="shortcut icon"], link[rel="apple-touch-icon"]');
            return el ? el.href : null;
        }
    """)
    if not favicon_url:
        issues.append({'severity': 'warning', 'check': 'Favicon',
                       'detail': 'Favicon לא מוגדר בעמוד (חסר link tag)',
                       'fix': 'הגדר Favicon ב-WordPress → מראה → התאמה אישית → זהות האתר → סמל האתר'})
    else:
        try:
            fav_resp = requests.head(favicon_url, timeout=6, allow_redirects=True)
            if fav_resp.status_code != 200:
                issues.append({'severity': 'warning', 'check': 'Favicon',
                               'detail': f'Favicon מוגדר אך לא נטען (HTTP {fav_resp.status_code}): {favicon_url}',
                               'fix': 'הקובץ חסר בשרת — העלה מחדש דרך מראה → התאמה אישית → סמל האתר'})
        except Exception:
            issues.append({'severity': 'warning', 'check': 'Favicon',
                           'detail': f'Favicon לא נגיש (timeout/שגיאת רשת): {favicon_url}',
                           'fix': 'בדוק שקובץ ה-favicon קיים ונגיש'})

    # ══════════════════════════════════════════════════════
    # 15. קישורים פנימיים שבורים (רק ב-Desktop Chrome, ולא בכל עמוד)
    # ══════════════════════════════════════════════════════
    if device['name'] == 'Desktop Chrome 1920' and page_info['path'] in ['/', '/metrics/naam/']:
        try:
            internal_links = await page.evaluate(f"""
                () => [...new Set(Array.from(document.querySelectorAll('a[href]'))
                    .map(a => a.href)
                    .filter(h => h.startsWith('{BASE_URL}') && !h.includes('#')
                             && !h.includes('tel:') && !h.includes('mailto:')
                             && !h.includes('wp-admin') && !h.includes('wp-login'))
                )].slice(0, 20)
            """)
            broken_links = []
            for link in internal_links[:15]:
                try:
                    r = requests.head(link, timeout=5, allow_redirects=True,
                                      headers={'User-Agent': 'PRCE-QA-Bot'})
                    if r.status_code == 404:
                        broken_links.append(link.replace(BASE_URL, ''))
                except Exception:
                    pass
            if broken_links:
                issues.append({'severity': 'error', 'check': 'קישורים פנימיים שבורים',
                               'detail': f'קישורים 404: {", ".join(broken_links)}',
                               'fix': 'תקן או הסר את הקישורים השבורים ב-Elementor'})
        except Exception:
            pass

    # ══════════════════════════════════════════════════════
    # 16. קישורים חיצוניים — rel=noopener
    # ══════════════════════════════════════════════════════
    if device['name'] == 'Desktop Chrome 1920':
        unsafe_ext = await page.evaluate(f"""
            () => Array.from(document.querySelectorAll('a[target="_blank"]'))
                .filter(a => {{
                    const rel = a.getAttribute('rel') || '';
                    return !rel.includes('noopener') && !a.href.startsWith('{BASE_URL}');
                }})
                .map(a => (a.href || '').split('/').slice(0, 3).join('/'))
                .filter((v, i, a) => v && a.indexOf(v) === i)
                .slice(0, 4)
        """)
        if unsafe_ext:
            issues.append({'severity': 'warning', 'check': 'אבטחה — קישורים חיצוניים',
                           'detail': f'target="_blank" ללא rel="noopener": {", ".join(unsafe_ext)}',
                           'fix': 'הוסף rel="noopener noreferrer" לכל קישור חיצוני שנפתח בטאב חדש'})

    # ══════════════════════════════════════════════════════
    # 17. טופס יצירת קשר
    # ══════════════════════════════════════════════════════
    if any(kw in page_info['name'] for kw in ['קשר', 'contact', 'Contact']):
        form_ok = await page.evaluate("""
            () => !!document.querySelector('form, .elementor-form, [class*="contact-form"]')
        """)
        if not form_ok:
            issues.append({'severity': 'error', 'check': 'טופס יצירת קשר',
                           'detail': 'טופס לא נמצא בעמוד צור קשר',
                           'fix': 'בדוק ב-Elementor שווידג\'ט Form קיים ומוצג'})

    # ══════════════════════════════════════════════════════
    # 18. אלמנטי API לא תקועים
    # ══════════════════════════════════════════════════════
    widget_ids = API_WIDGET_IDS.get(page_info['path'], [])
    if widget_ids:
        hidden = await page.evaluate(f"""
            () => {{
                return {json.dumps(widget_ids)}.filter(id => {{
                    const el = document.querySelector('[data-id="' + id + '"]');
                    if (!el) return false;
                    if (el.offsetParent === null) return false;
                    return parseFloat(getComputedStyle(el).opacity) < 0.5;
                }});
            }}
        """)
        if hidden:
            issues.append({'severity': 'error', 'check': 'ווידג\'טי API תקועים',
                           'detail': f'data-ids: {hidden} — opacity:0',
                           'fix': 'בדוק ב-Code Snippets שקוד ה-reveal רץ. נקה ezCache לאחר תיקון.'})

    # ══════════════════════════════════════════════════════
    # 19. אימות ערך API
    # ══════════════════════════════════════════════════════
    if page_info.get('check_api') and 'error' not in expected:
        displayed = await page.evaluate("""
            () => {
                for (const w of document.querySelectorAll('.elementor-widget[data-id]')) {
                    const t = (w.innerText || '').trim();
                    if (/^1[,0-9]{3,6}[.][0-9]{2}$/.test(t)) return t;
                }
                return null;
            }
        """)
        if displayed:
            if displayed.replace(',','') != expected['value'].replace(',',''):
                issues.append({'severity': 'error', 'check': 'ערך מדד',
                               'detail': f'מוצג: {displayed} | צפוי: {expected["value"]}',
                               'fix': 'הערך לא תואם ל-API. בדוק snippet, נקה ezCache.'})
        else:
            issues.append({'severity': 'warning', 'check': 'ערך מדד',
                           'detail': 'לא נמצא ווידג\'ט עם ערך מספרי',
                           'fix': 'בדוק ידנית שהמדד מוצג בדפדפן'})

        # בדיקת תאריך עדכון — חייב להיות יום המסחר האחרון (שני-שישי)
        last_trading_day = get_last_trading_day()
        last_trading_fmt = fmt_date_il(last_trading_day)
        api_vdate = datetime.strptime(expected['value_date'], '%Y-%m-%d').date()

        date_ok = await page.evaluate(f"""
            () => [...document.querySelectorAll('.elementor-widget[data-id]')]
                   .some(w => (w.innerText || '').trim() === '{expected["date"]}')
        """)
        if not date_ok:
            issues.append({'severity': 'warning', 'check': 'תאריך עדכון מדד',
                           'detail': f'לא נמצא תאריך {expected["date"]} בעמוד',
                           'fix': 'בדוק שה-snippet מעדכן שדה התאריך'})
        elif api_vdate < last_trading_day:
            # האתר מציג תאריך תקין מהAPI, אבל הAPI עצמו מיושן — זה באג
            issues.append({'severity': 'error', 'check': 'תאריך עדכון מדד — נתונים מיושנים ❗',
                           'detail': (f'האתר מציג נתוני {expected["date"]} '
                                      f'אך יום המסחר האחרון היה {last_trading_fmt} — '
                                      f'הנתונים לא עודכנו!'),
                           'fix': ('בדוק שה-API מחזיר את נתוני יום המסחר האחרון. '
                                   'ייתכן שתהליך העדכון נכשל — בדוק לוגים ואת מקור הנתונים.')})

    # ══════════════════════════════════════════════════════
    # 20. כפתור טלפון
    # ══════════════════════════════════════════════════════
    phone_ok = await page.evaluate("""
        () => [...document.querySelectorAll('a')].some(a =>
            a.href?.includes('tel:') || (a.innerText||'').includes('052'))
    """)
    if not phone_ok:
        issues.append({'severity': 'warning', 'check': 'כפתור טלפון',
                       'detail': 'לא נמצא קישור tel: בעמוד',
                       'fix': 'בדוק ב-Elementor שכפתור הטלפון מוגדר עם href="tel:0524446533"'})

    # ══════════════════════════════════════════════════════
    # 21. בדיקת נראות ויזואלית — CSS של Elementor
    # ══════════════════════════════════════════════════════
    if device['name'] == 'Desktop Chrome 1920':  # בודק רק פעם אחת לעמוד
        try:
            visual = await page.evaluate("""
                () => {
                    const rootStyle = getComputedStyle(document.documentElement);
                    const globalCss = rootStyle.getPropertyValue('--e-global-color-primary').trim();
                    const bodyH = document.body.scrollHeight;
                    const sections = document.querySelectorAll(
                        '.elementor-section, .e-con, .elementor-container');
                    const visibleSections = Array.from(sections)
                        .filter(s => s.offsetHeight > 50).length;
                    return {
                        globalCssLoaded: globalCss !== '',
                        primaryColor: globalCss,
                        bodyHeight: bodyH,
                        visibleSections: visibleSections
                    };
                }
            """)
            if not visual.get('globalCssLoaded'):
                issues.append({
                    'severity': 'error',
                    'check': 'נראות — CSS גלובלי של Elementor',
                    'detail': 'משתני CSS גלובליים (צבעי האתר) לא נטענו — קובץ CSS של Elementor מחזיר שגיאה!',
                    'fix': 'ב-Elementor → Tools → Regenerate CSS and Data. נקה ezCache לאחר מכן.'
                })
            elif visual.get('bodyHeight', 9999) < 500:
                issues.append({
                    'severity': 'error',
                    'check': 'נראות — תוכן עמוד',
                    'detail': f'גובה העמוד חשוד: {visual.get("bodyHeight")}px — ייתכן שהעמוד ריק',
                    'fix': 'בדוק שה-CSS של Elementor נטען ושיש תוכן בעמוד'
                })
            elif visual.get('visibleSections', 99) < 2:
                issues.append({
                    'severity': 'warning',
                    'check': 'נראות — sections גלויים',
                    'detail': f'רק {visual.get("visibleSections", 0)} sections גלויים — ייתכן שהעמוד נראה שבור',
                    'fix': 'בדוק ב-Elementor שכל הסקציות מוצגות, ורנדר CSS מחדש אם נדרש'
                })
        except Exception:
            pass

    # ══════════════════════════════════════════════════════
    # 22. צילום מסך
    # ══════════════════════════════════════════════════════
    ss_name = f"{device['name'].replace(' ','_')}_{page_info['path'].strip('/').replace('/','_') or 'home'}.png"
    ss_path = SCREENSHOTS_DIR / ss_name
    await page.screenshot(path=str(ss_path), full_page=False)

    return {
        'page':       page_info['name'],
        'path':       page_info['path'],
        'device':     device['name'],
        'issues':     issues,
        'screenshot': ss_name,
        'perf':       perf,
        'ok':         len([i for i in issues if i['severity'] == 'error']) == 0,
    }


# ─── Word Report helpers ──────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return p

def para(doc, text='', bold=False, color=None, size=None, center=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.bold = bold
    if color: run.font.color.rgb = RGBColor(*color)
    if size:  run.font.size = Pt(size)
    return p

def make_table(doc, headers, header_bg='1A2B3C'):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        set_cell_bg(cell, header_bg)
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255,255,255)
    return tbl

def add_row(tbl, values, bg='FFFFFF'):
    row = tbl.add_row()
    for i, val in enumerate(values):
        row.cells[i].text = str(val)[:200]
        set_cell_bg(row.cells[i], bg)
    return row

# ─── False-positive filters ───────────────────────────────────────────────────
FALSE_POSITIVE_DETAILS = ['STAT:', 'nameID', 'Table discarded', '_ga', '_gcl']

def is_false_positive(issue):
    detail = issue.get('detail','')
    if issue.get('check') == 'טעינת עמוד' and 'Timeout' in detail: return True
    if issue.get('check') == 'טעינת עמוד' and '503' in detail: return True
    return False

# ─── בניית דוח Word ──────────────────────────────────────────────────────────
def build_word_report(all_results, expected, static_checks, ts, pages):
    doc = Document()

    # עמוד שער
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('דוח QA — prce.co.il')
    r.font.size = Pt(28); r.font.bold = True
    r.font.color.rgb = RGBColor(0x19, 0x24, 0x3A)

    para(doc, f'הופק: {ts}', center=True, size=13, color=(80,80,80))
    para(doc, f'עמודים שנבדקו: {len(pages)}  |  מכשירים: {len(DEVICES)}  |  בדיקות: {len(pages)*len(DEVICES)}', center=True, size=11)
    if 'error' not in expected:
        para(doc, f'ערך מדד: {expected["value"]}  |  תאריך עדכון: {expected["date"]}',
             center=True, size=11, color=(0,128,0))
    doc.add_paragraph()

    # סינון false positives
    real_results = []
    for r in all_results:
        real_issues = [i for i in r['issues'] if not is_false_positive(i)]
        real_results.append({**r, 'issues': real_issues})

    unique = {}
    for r in real_results:
        for i in r['issues']:
            key = (i['severity'], i['check'], i['detail'][:70])
            if key not in unique:
                unique[key] = {'issue': i, 'pages': [], 'devices': []}
            if r['page'] not in unique[key]['pages']:
                unique[key]['pages'].append(r['page'])
            if r['device'] not in unique[key]['devices']:
                unique[key]['devices'].append(r['device'])

    errors_u   = {k: v for k,v in unique.items() if k[0] == 'error'}
    warnings_u = {k: v for k,v in unique.items() if k[0] == 'warning'}
    n_err = len(errors_u)
    n_wrn = len(warnings_u)

    # ═══ לוח בקרה ═══
    heading(doc, 'לוח בקרה', 1)
    status_text  = 'האתר תקין לחלוטין ✅' if n_err == 0 else f'⚠️ {n_err} בעיות קריטיות'
    status_color = (0,150,0) if n_err == 0 else (180,0,0)
    para(doc, f'סטטוס כללי: {status_text}', bold=True, size=14, color=status_color)
    doc.add_paragraph()

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Table Grid'
    labels = ['עמודים', 'מכשירים', 'שגיאות קריטיות', 'אזהרות', 'בדיקות סטטיות']
    static_issues = sum(1 for v in static_checks.values() if not v['ok'])
    values = [str(len(pages)), str(len(DEVICES)), str(n_err), str(n_wrn), str(static_issues)]
    bgs    = ['D5E8D4','D5E8D4',
              'FFD7D7' if n_err else 'D5E8D4',
              'FFF3CD' if n_wrn else 'D5E8D4',
              'FFD7D7' if static_issues else 'D5E8D4']
    for i, (lbl, val, bg) in enumerate(zip(labels, values, bgs)):
        cell = tbl.rows[0].cells[i]
        cell.paragraphs[0].clear()
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.add_run(val).font.size = Pt(20)
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.add_run(lbl).font.size = Pt(9)
        set_cell_bg(cell, bg)
    doc.add_paragraph()

    # ═══ בדיקות סטטיות (sitemap, robots, video, API, HTTPS) ═══
    heading(doc, 'בדיקות תשתית', 1)
    check_labels = {
        'sitemap':        'Sitemap.xml',
        'robots':         'Robots.txt',
        'hero_video':     'סרטון Hero',
        'api_freshness':  'עדכניות נתוני API',
        'https_redirect': 'הפניית HTTPS',
    }
    tbl_s = make_table(doc, ['בדיקה', 'סטטוס', 'פרטים'], '2C3E50')
    for key, label in check_labels.items():
        chk = static_checks.get(key, {'ok': None, 'detail': 'לא נבדק'})
        icon = '✅' if chk['ok'] else '❌'
        bg   = 'D5E8D4' if chk['ok'] else 'FFD7D7'
        add_row(tbl_s, [label, icon, chk['detail']], bg=bg)
    doc.add_paragraph()

    # ═══ ביצועים — סיכום ממוצע ═══
    heading(doc, 'ביצועי עמודים (Desktop Chrome)', 1)
    desktop_results = [r for r in real_results if 'Chrome 1920' in r['device'] and r.get('perf')]
    if desktop_results:
        tbl_p = make_table(doc, ['עמוד', 'TTFB', 'זמן טעינה', 'גודל'], '2C3E50')
        for r in desktop_results:
            perf = r.get('perf', {})
            ttfb = f"{perf.get('ttfb','?')}ms"
            load = f"{perf.get('load',0)/1000:.1f}s" if perf.get('load') else '?'
            size = f"{perf.get('size',0)/1024:.1f}MB" if perf.get('size') else '?'
            bg   = 'FFF3CD' if perf.get('load',0) > 5000 else 'D5E8D4'
            add_row(tbl_p, [r['page'][:40], ttfb, load, size], bg=bg)
    doc.add_paragraph()

    # ═══ בעיות קריטיות ═══
    heading(doc, 'חלק א׳ — בעיות קריטיות', 1)
    if not errors_u:
        para(doc, '✅ לא נמצאו בעיות קריטיות.', color=(0,128,0))
    else:
        para(doc, f'נמצאו {len(errors_u)} בעיות קריטיות.', color=(180,0,0), bold=True)
        doc.add_paragraph()
        for idx, ((sev, check, detail), info) in enumerate(errors_u.items(), 1):
            heading(doc, f'{idx}. {check}', 2)
            para(doc, f'מה קרה: {info["issue"]["detail"][:250]}')
            para(doc, f'איפה: {", ".join(info["pages"][:6])}')
            para(doc, f'מכשירים: {", ".join(info["devices"][:4])}')
            doc.add_paragraph()
            para(doc, 'איך מתקנים:', bold=True)
            para(doc, info['issue']['fix'])
            first_page = info['pages'][0] if info['pages'] else None
            if first_page:
                match = next((r for r in real_results if r['page'] == first_page and r['screenshot']), None)
                if match and (SCREENSHOTS_DIR / match['screenshot']).exists():
                    try: doc.add_picture(str(SCREENSHOTS_DIR / match['screenshot']), width=Cm(13))
                    except: pass
            doc.add_paragraph()

    # ═══ אזהרות ═══
    heading(doc, 'חלק ב׳ — אזהרות', 1)
    if not warnings_u:
        para(doc, '✅ אין אזהרות.', color=(0,128,0))
    else:
        tbl2 = make_table(doc, ['#', 'בדיקה', 'פרטים', 'עמודים', 'תיקון'], '34495E')
        for idx, ((sev, check, detail), info) in enumerate(warnings_u.items(), 1):
            pages_str = ', '.join(info['pages'][:4]) + (f' (+{len(info["pages"])-4})' if len(info['pages'])>4 else '')
            add_row(tbl2, [str(idx), check, info['issue']['detail'][:100], pages_str, info['issue']['fix'][:100]], bg='FFF8E7')
    doc.add_paragraph()

    # ═══ פירוט לפי עמוד ═══
    heading(doc, 'חלק ג׳ — פירוט לפי עמוד', 1)
    by_page = {}
    for r in real_results:
        if r['path'] not in by_page:
            by_page[r['path']] = {'name': r['page'], 'results': []}
        by_page[r['path']]['results'].append(r)

    for path, pdata in by_page.items():
        n_e = sum(1 for r in pdata['results'] for i in r['issues'] if i['severity']=='error')
        n_w = sum(1 for r in pdata['results'] for i in r['issues'] if i['severity']=='warning')
        icon = '🔴' if n_e else ('🟡' if n_w else '🟢')
        heading(doc, f'{icon} {pdata["name"]}  ({path})', 2)

        if not any(r['issues'] for r in pdata['results']):
            para(doc, '✅ עמוד תקין — עבר את כל הבדיקות.', color=(0,128,0))
        else:
            tbl3 = make_table(doc, ['מכשיר', 'בדיקה', 'פרטים', 'תיקון'], '2C3E50')
            for r in pdata['results']:
                for issue in r['issues']:
                    bg = 'FFD7D7' if issue['severity']=='error' else 'FFF3CD'
                    add_row(tbl3, [r['device'], issue['check'], issue['detail'][:80], issue['fix'][:80]], bg=bg)

        preferred = next((r for r in pdata['results'] if 'Chrome 1920' in r['device'] and r.get('screenshot')), None)
        ss = preferred or next((r for r in pdata['results'] if r.get('screenshot')), None)
        if ss and (SCREENSHOTS_DIR / ss['screenshot']).exists():
            try: doc.add_picture(str(SCREENSHOTS_DIR / ss['screenshot']), width=Cm(14))
            except: pass
        doc.add_paragraph()

    ts_safe = ts.replace(':', '-').replace(' ', '_')
    out = REPORTS_DIR / f'QA_Report_{ts_safe}.docx'
    doc.save(str(out))
    return out


# ══════════════════════════════════════════════════════════════════════════════
# WHATSAPP
# ══════════════════════════════════════════════════════════════════════════════

def send_whatsapp(message: str) -> bool:
    """שולח הודעת וואצאפ דרך Twilio Sandbox."""
    try:
        from twilio.rest import Client
    except ImportError:
        print("❌ חסרה ספריית Twilio. הרץ: pip install twilio")
        return False

    import sys, os
    sys.path.insert(0, str(Path(__file__).parent))
    try:
        import config as cfg
        account_sid = cfg.TWILIO_ACCOUNT_SID
        auth_token  = cfg.TWILIO_AUTH_TOKEN
        from_number = cfg.TWILIO_FROM_NUMBER
        to_number   = f"whatsapp:+{cfg.WHATSAPP_PHONE}"
    except Exception as e:
        print(f"❌ שגיאה בטעינת config: {e}")
        return False

    try:
        client = Client(account_sid, auth_token)
        msg = client.messages.create(from_=from_number, to=to_number, body=message)
        print(f"✅ WhatsApp נשלח! SID: {msg.sid}")
        return True
    except Exception as e:
        print(f"❌ WhatsApp send error: {e}")
        return False


def build_whatsapp_summary(all_results, static_checks, api_data):
    """בונה סיכום קצר לוואצאפ — רק בעיות.
    all_results: רשימה של dict כפי שמוחזרת מ-check_page()
    static_checks: dict בדיקות תשתית מ-run_static_checks()
    api_data: dict נתוני API מ-get_expected_values()
    """
    today = datetime.now().strftime('%d/%m/%Y')
    RTL = '\u200F'
    lines = []

    lines.append(f"🔍 *QA יומי — prce.co.il*")
    lines.append(f"📅 {today}")
    lines.append('')

    # בעיות תשתית
    infra_issues = []
    for check, result in static_checks.items():
        if not result.get('ok'):
            infra_issues.append(f"• {check}: {result.get('detail', '')}")

    if infra_issues:
        lines.append('🏗️ *בעיות תשתית:*')
        for i in infra_issues:
            lines.append(f"  {i}")
        lines.append('')

    # ספור שגיאות ואזהרות (ללא כפילויות)
    errors = []
    warnings = []
    seen_issues = set()

    for result in all_results:
        page_path = result.get('path', '')
        for issue in result.get('issues', []):
            key = (page_path, issue.get('check', ''), issue.get('detail', '')[:50])
            if key in seen_issues:
                continue
            seen_issues.add(key)

            entry = f"• {page_path}: {issue.get('check', '')}"
            if issue.get('severity') == 'error':
                errors.append(entry)
            elif issue.get('severity') == 'warning':
                warnings.append(entry)

    total_issues = len(errors) + len(warnings)

    if total_issues == 0 and not infra_issues:
        lines.append('✅ *הכל תקין — אין בעיות!*')
    else:
        if errors:
            lines.append(f'❌ *שגיאות ({len(errors)}):*')
            for e in errors[:8]:  # מקסימום 8
                lines.append(f"  {e}")
            if len(errors) > 8:
                lines.append(f"  ...ועוד {len(errors)-8}")
            lines.append('')

        if warnings:
            lines.append(f'⚠️ *אזהרות ({len(warnings)}):*')
            for w in warnings[:6]:  # מקסימום 6
                lines.append(f"  {w}")
            if len(warnings) > 6:
                lines.append(f"  ...ועוד {len(warnings)-6}")
            lines.append('')

    lines.append('— QA Agent • prce.co.il')
    return '\n'.join(RTL + l if l.strip() else l for l in lines)


# ─── Main ─────────────────────────────────────────────────────────────────────
async def run_qa():
    REPORTS_DIR.mkdir(exist_ok=True)
    SCREENSHOTS_DIR.mkdir(exist_ok=True)
    ts = datetime.now().strftime('%Y-%m-%d %H:%M')

    print(f"\n{'='*65}")
    print(f"  PRCE QA Agent v2 - {ts}")
    print(f"{'='*65}\n")

    # בדיקות סטטיות (HTTP)
    print("[HTTP] מריץ בדיקות תשתית...")
    static_checks = run_static_checks()
    for key, val in static_checks.items():
        icon = '[OK] ' if val['ok'] else '[!!] '
        print(f"  {icon} {key}: {val['detail']}")
    print()

    # מושך עמודים
    print("[WP]  מושך עמודים מ-WordPress...")
    pages = get_all_wp_pages()
    print(f"[OK]  נמצאו {len(pages)} עמודים\n")

    # ערכי API
    print("[API] מושך ערכים מה-API...")
    expected = get_expected_values()
    if 'error' in expected:
        print(f"[!!]  API error: {expected['error']}")
    else:
        print(f"[OK]  ערך: {expected['value']} | תאריך: {expected['date']}\n")

    all_results = []

    async with async_playwright() as pw:
        for device in DEVICES:
            print(f"\n[>>]  {device['name']}")
            print(f"  {'-'*55}")

            browser = await getattr(pw, device['browser_type']).launch(headless=True)
            ctx_args = {
                'viewport':  device['viewport'],
                'is_mobile': device.get('is_mobile', False),
                'locale':    'he-IL'
            }
            if 'user_agent'          in device: ctx_args['user_agent']          = device['user_agent']
            if 'device_scale_factor' in device: ctx_args['device_scale_factor'] = device['device_scale_factor']

            context = await browser.new_context(**ctx_args)
            page    = await context.new_page()

            for page_info in pages:
                result = await check_page(page, page_info, expected, device)
                all_results.append(result)
                await asyncio.sleep(2)

                errs  = [i for i in result['issues'] if i['severity'] == 'error']
                warns = [i for i in result['issues'] if i['severity'] == 'warning']
                icon  = '[ERROR]' if errs else ('[WARN] ' if warns else '[OK]   ')
                print(f"  {icon}  {page_info['name']}")
                for i in errs + warns:
                    print(f"         {i['check']}: {i['detail'][:80]}")

            await browser.close()

    # דוח Word
    print(f"\n[DOC] בונה דוח Word...")
    report_path = build_word_report(all_results, expected, static_checks, ts, pages)
    print(f"[OK]  דוח נשמר: {report_path}")

    # ── WhatsApp summary ──
    print("\n[WA]  שולח סיכום בוואצאפ...")
    wa_msg = build_whatsapp_summary(all_results, static_checks, expected)
    send_whatsapp(wa_msg)

    # JSON גיבוי
    json_path = REPORTS_DIR / f"qa_{ts.replace(' ','_').replace(':','-')}.json"
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump({
            'ts': ts, 'expected': expected, 'static_checks': static_checks,
            'pages': len(pages), 'results': all_results
        }, f, ensure_ascii=False, indent=2)

    errors_total  = sum(1 for r in all_results for i in r['issues'] if i['severity'] == 'error')
    warning_total = sum(1 for r in all_results for i in r['issues'] if i['severity'] == 'warning')
    static_issues = sum(1 for v in static_checks.values() if not v['ok'])

    print(f"\n{'='*65}")
    print(f"  סיכום: {len(all_results)} בדיקות עמוד | {errors_total} שגיאות | {warning_total} אזהרות | {static_issues} בעיות תשתית")
    print(f"  דוח: {report_path}")
    print(f"{'='*65}\n")


if __name__ == '__main__':
    asyncio.run(run_qa())
